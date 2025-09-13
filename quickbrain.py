import streamlit as st
import fitz 
from langchain_huggingface import HuggingFaceEmbeddings
import hashlib
import re
from pinecone import Pinecone
from langchain_pinecone import PineconeVectorStore
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain_community.llms import SambaNovaCloud
from langchain.chains import LLMChain
from langchain.memory import ConversationBufferMemory
from langchain.prompts import PromptTemplate
from langchain_community.utilities import GoogleSerperAPIWrapper
from supabase import create_client, Client
from collections import deque
from dotenv import load_dotenv
import os
import time
import pytz
import hashlib
from datetime import datetime
import uuid

load_dotenv(dotenv_path="../crewai/.env")
supabase_url = os.getenv("SUPABASE_URL")
supabase_key = os.getenv("SUPABASE_KEY")

supabase: Client = None
if supabase_url and supabase_key:
    supabase = create_client(supabase_url, supabase_key)
else:
    st.error("Supabase credentials not found in environment variables.")

st.title("QuickBrain")

def now_iso():
    tz = pytz.timezone('Asia/Kolkata')
    return datetime.now(tz).isoformat(timespec="seconds")

def extract_text_from_file(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        text = ""
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    elif name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")
    elif name.endswith(".docx"):
        from docx import Document
        doc = Document(uploaded_file)
        return "\n".join(para.text for para in doc.paragraphs)
    else:
        raise ValueError("Unsupported file type")

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def ensure_memory_from_chat(chat):
    mem = ConversationBufferMemory(
        memory_key="chat_history",
        return_messages=True,
        output_key="answer",
    )
    for m in chat.get("messages", []):
        if m["role"] == "user":
            mem.chat_memory.add_user_message(m["content"])
        elif m["role"] == "assistant":
            mem.chat_memory.add_ai_message(m["content"])
    return mem

def render_history_text(mem: ConversationBufferMemory) -> str:
    vars_ = mem.load_memory_variables({})
    msgs = vars_.get("chat_history", [])
    parts = []
    for m in msgs:
        t = getattr(m, "type", None) or m.__class__.__name__.lower()
        if t in ("human", "humanmessage"):
            parts.append(f"User: {m.content}")
        elif t in ("ai", "aimessage"):
            parts.append(f"Assistant: {m.content}")
        else:
            parts.append(m.content)
    return "\n".join(parts)

def _clear_chat_search():
    st.session_state.chat_search_query = ""
    
def _toggle_web_sidebar():
    st.session_state.use_web_search = not st.session_state.use_web_search
    
def _begin_processing():
    st.session_state.is_processing_docs = True

def allow(action: str, limit: int, window_sec: int) -> bool:
    key = f"rl_{action}"
    if key not in st.session_state:
        st.session_state[key] = deque()
    q = st.session_state[key]
    now = time.time()
    while q and now - q[0] > window_sec:
        q.popleft()
    if len(q) < limit:
        q.append(now)
        return True
    return False

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "show_signup" not in st.session_state:
    st.session_state.show_signup = False
if "show_login" not in st.session_state:
    st.session_state.show_login = True
if "selected_chat_id" not in st.session_state:
    st.session_state.selected_chat_id = None
if "chats" not in st.session_state:
    st.session_state.chats = {}
if "memory" not in st.session_state:
    st.session_state.memory = None
if "is_processing_docs" not in st.session_state:
    st.session_state.is_processing_docs = False
if "use_web_search" not in st.session_state:
    st.session_state.use_web_search = False
if "chat_search_query" not in st.session_state:
    st.session_state.chat_search_query = ""

if not st.session_state.logged_in:
    if st.session_state.show_signup:
        st.subheader("Signup")
        with st.form("signup_form"):
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            submit = st.form_submit_button("Signup")
            if submit:
                if not (username and password):
                    st.error("Please fill in all fields.")
                elif not re.fullmatch(r"^[A-Za-z0-9._%+-]+@gmail\.com$", username):
                    st.error("Username must be a valid address ending with Gmail Account.")
                elif supabase is None:
                    st.error("Supabase is not configured.")
                else:
                    existing_user = supabase.table("profiles").select("*").eq("username", username).execute()
                    if existing_user.data:
                        st.error("Username already exists.")
                    else:
                        password_hash = hash_password(password)
                        resp = supabase.table("profiles").insert(
                            {"username": username, "password": password_hash, "chats": []}
                        ).execute()

                        user_row = None
                        try:
                            if resp and getattr(resp, "data", None):
                                user_row = resp.data[0]
                            else:
                                fetched = supabase.table("profiles").select("*").eq("username", username).execute()
                                user_row = fetched.data[0] if fetched.data else None
                        except Exception:
                            user_row = None

                        if not user_row:
                            st.success("Signup successful. Please login.")
                            st.session_state.show_signup = False
                            st.session_state.show_login = True
                            st.rerun()

                        st.session_state.logged_in = True
                        st.session_state.user_id = user_row["id"]
                        st.session_state.username = username

                        st.session_state.chats = {}
                        new_chat_id = str(uuid.uuid4())
                        ts = now_iso()
                        new_chat = {
                            "id": new_chat_id,
                            "title": "New Chat",
                            "messages": [],
                            "created_at": ts,
                            "updated_at": ts,
                        }
                        st.session_state.chats[new_chat_id] = new_chat
                        st.session_state.selected_chat_id = new_chat_id
                        st.session_state.memory = ensure_memory_from_chat(new_chat)

                        try:
                            supabase.table("profiles").update(
                                {"chats": list(st.session_state.chats.values())}
                            ).eq("id", st.session_state.user_id).execute()
                        except Exception as e:
                            st.error(f"Error saving chats: {str(e)}")

                        st.session_state.show_signup = False
                        st.session_state.show_login = False
                        st.success("Signup successful. Welcome!")
                        st.rerun()
        if st.button("Login", key="login_acct_btn"):
            st.session_state.show_login = True
            st.session_state.show_signup = False
            st.rerun()
                

    if st.session_state.show_login:
        st.subheader("Login")
        with st.form("login_form"):
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            submit = st.form_submit_button("Login")
            if submit:
                if not (username and password):
                    st.error("Please fill in all fields.")
                elif supabase is None:
                    st.error("Supabase is not configured.")
                else:
                    user = supabase.table("profiles").select("*").eq("username", username).execute()
                    if user.data and hash_password(password) == user.data[0]["password"]:
                        st.session_state.logged_in = True
                        st.session_state.user_id = user.data[0]["id"]
                        st.session_state.username = username

                        chats = user.data[0].get("chats", [])
                        for c in chats:
                            if "created_at" not in c:
                                c["created_at"] = c.get("date", now_iso())
                            if "updated_at" not in c:
                                c["updated_at"] = c.get("created_at", now_iso())
                        st.session_state.chats = {chat["id"]: chat for chat in chats}

                        new_chat_id = str(uuid.uuid4())
                        ts = now_iso()
                        new_chat = {
                            "id": new_chat_id,
                            "title": "New Chat",
                            "messages": [],
                            "created_at": ts,
                            "updated_at": ts,
                        }
                        st.session_state.chats[new_chat_id] = new_chat
                        st.session_state.selected_chat_id = new_chat_id
                        st.session_state.memory = ensure_memory_from_chat(new_chat)

                        try:
                            supabase.table("profiles").update(
                                {"chats": list(st.session_state.chats.values())}
                            ).eq("id", st.session_state.user_id).execute()
                        except Exception as e:
                            st.error(f"Error saving chats: {str(e)}")

                        st.session_state.show_login = False
                        st.success("Login successful.")
                        st.rerun()
                    else:
                        st.error("Invalid username or password.")
        if st.button("Create new account", key="create_acct_btn"):
            st.session_state.show_login = False
            st.session_state.show_signup = True
            st.rerun()

    st.stop()

st.sidebar.title("Settings")
api_key = st.sidebar.text_input(
    "SambaNova API Key",
    type="password",
    disabled=st.session_state.is_processing_docs
)
st.sidebar.caption("Don’t have one? Get it at [sambanova.ai](https://sambanova.ai/).")

model_choice = st.sidebar.selectbox(
    "Choose model",
    options=["Llama-3.3-Swallow-70B-Instruct-v0.4", "Llama-4-Maverick-17B-128E-Instruct", "Meta-Llama-3.1-8B-Instruct", "Meta-Llama-3.3-70B-Instruct"],
    index=3,
    disabled=st.session_state.is_processing_docs
)

serper_api_key = st.sidebar.text_input(
    "Serper API Key",
    type="password",
    disabled=st.session_state.is_processing_docs
)
st.sidebar.caption("Don’t have one? Get it at [serper.dev](https://serper.dev/).")

ws_label = f"Web search: {'On' if st.session_state.use_web_search and serper_api_key else 'Off'}"
st.sidebar.button(ws_label, key="web_toggle_sidebar", on_click=_toggle_web_sidebar, disabled=st.session_state.is_processing_docs)

st.success(f"Logged in as {st.session_state.username}")
if st.sidebar.button("Logout", key="logout_btn", disabled=st.session_state.is_processing_docs):
    st.session_state.logged_in = False
    st.session_state.chats = {}
    st.session_state.memory = None
    st.session_state.selected_chat_id = None
    st.session_state.show_login = True
    st.session_state.show_signup = False
    st.rerun()

MAX_FILES = 3
uploaded_files = st.file_uploader(
    "Upload up to 3 files each under 500KB",
    type=["pdf", "txt", "docx"],
    accept_multiple_files=True,
    disabled=st.session_state.is_processing_docs,
    key="files",
)

uploaded_files = uploaded_files or []
if len(uploaded_files) > MAX_FILES:
    st.error("Select at most 3 Files.")
    st.stop()

MAX_FILE_SIZE_BYTES = 500 * 1024
oversized_files = [uf for uf in uploaded_files if getattr(uf, "size", 0) > MAX_FILE_SIZE_BYTES]
if oversized_files:
    names = ", ".join(f.name for f in oversized_files)
    st.error(f"Each file must be under 500 KB. Too large: {names}")
    st.stop()


st.button(
    "Process Documents",
    key="process_docs_btn",
    on_click=_begin_processing,
    disabled=st.session_state.is_processing_docs or not (uploaded_files and api_key),
)

if st.session_state.is_processing_docs:
    with st.spinner("Processing documents..."):
        all_docs = []
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        for uf in uploaded_files or []:
            text = extract_text_from_file(uf)
            for chunk in splitter.split_text(text):
                all_docs.append(Document(page_content=chunk, metadata={"source": uf.name}))
        
        embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-large-en-v1.5")
        pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
        index = pc.Index(os.getenv('PINECONE_INDEX_NAME'))
        vectorstore = PineconeVectorStore(
            index=index,
            embedding=embeddings,
            namespace=st.session_state.username, 
        )
        if all_docs:
            ids = [
                hashlib.sha1((d.page_content + d.metadata.get("source", "")).encode()).hexdigest()
                for d in all_docs
            ]
            vectorstore.add_documents(all_docs, ids=ids)

        st.session_state.vectorstore = vectorstore
    st.session_state.is_processing_docs = False
    st.success("Documents processed and indexed.")
    st.rerun()
    
st.sidebar.title("Search Chats")
with st.sidebar:
    c1, c2 = st.columns([0.85, 0.15])
    with c1:
        q = st.text_input(
            "Search chats",
            value=st.session_state.chat_search_query,
            label_visibility="collapsed",
            placeholder="Search conversations…",
            key="chat_search_input",
            disabled=st.session_state.is_processing_docs,
        )
        
        if q != st.session_state.chat_search_query:
            st.session_state.chat_search_query = q
    with c2:
        if st.button("✕", key="chat_search_clear", help="Clear search", disabled=st.session_state.is_processing_docs):
            _clear_chat_search()
            st.rerun()

st.sidebar.title("Chats")
if not st.session_state.is_processing_docs:
    disable_new_chat = False
    if st.session_state.selected_chat_id:
        current_chat = st.session_state.chats[st.session_state.selected_chat_id]
        if current_chat["title"] == "New Chat" and len(current_chat.get("messages", [])) == 0:
            disable_new_chat = True

    if st.sidebar.button("New Chat", key="new_chat_btn", disabled=disable_new_chat):
        new_chat_id = str(uuid.uuid4())
        ts = now_iso()
        new_chat = {
            "id": new_chat_id,
            "title": "New Chat",
            "messages": [],
            "created_at": ts,
            "updated_at": ts,
        }
        st.session_state.chats[new_chat_id] = new_chat
        st.session_state.selected_chat_id = new_chat_id
        st.session_state.memory = ensure_memory_from_chat(new_chat)
        if supabase is not None:
            try:
                supabase.table("profiles").update(
                    {"chats": list(st.session_state.chats.values())}
                ).eq("id", st.session_state.user_id).execute()
            except Exception as e:
                st.error(f"Error saving chats: {str(e)}")
        st.rerun()

    sorted_chats = sorted(
        st.session_state.chats.values(),
        key=lambda x: x.get("updated_at", x.get("created_at", "1970-01-01T00:00:00")),
        reverse=True,
    )
    chat_search_q = st.session_state.chat_search_query.strip().lower()
    if chat_search_q:
        sorted_chats = [
            c for c in sorted_chats
            if any(
                chat_search_q in (m.get("content", "").lower())
                for m in c.get("messages", [])
            )
        ]
    for chat in sorted_chats:
        ts = chat.get("updated_at", chat.get("created_at", ""))[:16].replace("T", " ")
        label = f"{ts} - {chat.get('title','New Chat')}"
        if st.sidebar.button(label, key=f"open_chat_{chat['id']}"):
            st.session_state.selected_chat_id = chat["id"]
            st.session_state.memory = ensure_memory_from_chat(chat)
            st.rerun()
else:
    st.sidebar.info("Processing documents. Chats hidden.")

if st.session_state.selected_chat_id:
    current_chat = st.session_state.chats[st.session_state.selected_chat_id]
    if st.session_state.memory is None:
        st.session_state.memory = ensure_memory_from_chat(current_chat)

    for message in current_chat.get("messages", []):
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
else:
    st.info("Select a chat or start a new one.")

if api_key and st.session_state.selected_chat_id and not st.session_state.is_processing_docs:
    if st.session_state.use_web_search and not serper_api_key:
        st.warning("No Serper API key is set. Falling back to LLM.")

    query = st.chat_input("Your question")
    if query:
        current_chat = st.session_state.chats[st.session_state.selected_chat_id]
        current_time = now_iso()

        if current_chat["title"] == "New Chat":
            new_title = query[:50] + "..." if len(query) > 50 else query
            current_chat["title"] = new_title
            st.session_state["_sidebar_title_just_updated"] = True

        current_chat["messages"].append({"role": "user", "content": query, "date": current_time})
        with st.chat_message("user"):
            st.markdown(query)

        with st.spinner("Generating answer..."):
            if not allow("llm_calls", limit=50, window_sec=180):
                st.warning("Rate limit: 50 requests per minute.")
                st.stop()
            llm = SambaNovaCloud(model=model_choice, sambanova_api_key=api_key, temperature=0.4, top_p=0.8, top_k=1000)

            history_text = render_history_text(st.session_state.memory)

            if st.session_state.use_web_search and serper_api_key:
                search_tool = GoogleSerperAPIWrapper(serper_api_key=serper_api_key, k=5)
                sr = search_tool.results(query)
                organic = sr.get("organic", [])
                search_results = "\n".join(
                    f"Snippet: {o.get('snippet','')}\nLink: {o.get('link','')}" for o in organic
                )
                web_prompt = PromptTemplate.from_template(
                    "Based on the following web search results, answer concisely and accurately. "
                    "Include key sources.\n\nSearch Results:\n{search_results}\n\nQuestion: {question}\n\nAnswer:"
                )
                web_chain = LLMChain(llm=llm, prompt=web_prompt, output_key="ans")
                web = web_chain({"search_results": search_results, "question": query})
                final_answer = web["ans"].strip()
                sources = [o.get("link","") for o in organic if o.get("link")]
                if sources:
                    source_block = "\n\n**Web Sources used:**\n" + "\n".join(f"- {s}" for s in sources)
                else:
                    source_block = ""
            else:
                probe_prompt = PromptTemplate.from_template(
                    "Use chat history as context only if you can answer this question. Your answers should be non-deterministic and no one should be able to guess the answer."
                    "If the history lacks enough info, reply exactly: NEEDS_EXTERNAL\n\n"
                    "Chat History:\n{chat_history}\n\nQuestion: {question}\n\nAnswer:"
                )
                probe_chain = LLMChain(llm=llm, prompt=probe_prompt, output_key="answer")
                probe = probe_chain({"chat_history": history_text, "question": query})
                probe_answer = probe["answer"].strip()

                if probe_answer != "NEEDS_EXTERNAL":
                    final_answer = probe_answer
                    source_block = ""
                else:
                    final_answer, source_block = None, ""
                    rag_answer, rag_sources = None, None
                    if "vectorstore" in st.session_state:
                        retriever = st.session_state.vectorstore.as_retriever(
                            search_kwargs={"k": 10}, search_type="mmr"
                        )
                        docs = retriever.get_relevant_documents(query)
                        context = "\n\n".join(d.page_content for d in docs)
                        combine_prompt = PromptTemplate.from_template(
                            "Use the context to answer the question. "
                            "If insufficient, say 'Information not found in the provided documents.'\n\n"
                            "{context}\n\nQuestion: {question}\n\nAnswer:"
                        )
                        combine_chain = LLMChain(llm=llm, prompt=combine_prompt, output_key="ans")
                        rag = combine_chain({"context": context, "question": query})
                        rag_answer = rag["ans"].strip()
                        rag_sources = sorted({d.metadata.get("source", "") for d in docs if d and d.metadata.get("source")})

                    if rag_answer and "Information not found in the provided documents." not in rag_answer:
                        final_answer = rag_answer
                        if rag_sources:
                            source_block = "\n\n**Sources used:**\n" + "\n".join(f"- {s}" for s in rag_sources)
                    else:
                        direct_prompt = PromptTemplate.from_template(
                            "Answer clearly and concisely.\n\nQuestion: {q}\n\nAnswer:"
                        )
                        direct_chain = LLMChain(llm=llm, prompt=direct_prompt, output_key="ans")
                        direct = direct_chain({"q": query})
                        final_answer = direct["ans"].strip()

            eval_prompt = PromptTemplate.from_template(
                "Rate the answer on a scale of 1-10 for accuracy, relevance, and completeness. "
                "Start with 'Score: X/10' where X is the floating point number, then explain briefly.\n\n"
                "Question: {question}\nAnswer: {answer}\n\nEvaluation:"
            )
            eval_chain = LLMChain(llm=llm, prompt=eval_prompt, output_key="eval")
            evaluation = eval_chain({"question": query, "answer": final_answer})["eval"]

            try:
                score_line = evaluation.split('\n')[0]
                score_str = score_line.split(':')[1].split('/')[0].strip()
                score = int(score_str)
            except:
                match = re.search(r'\b([1-9]|10)\b', evaluation)
                score = int(match.group(1)) if match else 2 # Default to 5 if parsing fails
            assistant_response = f"**Answer:**\n{final_answer}{source_block}"

        current_chat["messages"].append({"role": "assistant", "content": assistant_response, "score": score,"date": current_time})
        current_chat["updated_at"] = now_iso()

        st.session_state.memory.chat_memory.add_user_message(query)
        st.session_state.memory.chat_memory.add_ai_message(assistant_response)

        with st.chat_message("assistant"):
            st.markdown(assistant_response)

        if supabase is not None:
            try:
                supabase.table("profiles").update(
                    {"chats": list(st.session_state.chats.values())}
                ).eq("id", st.session_state.user_id).execute()
            except Exception as e:
                st.error(f"Error saving chats: {str(e)}")

        if st.session_state.pop("_sidebar_title_just_updated", False):
            st.rerun()
else:
    if not api_key:
        st.info("Enter your SambaNova API key in the sidebar.")
    elif st.session_state.selected_chat_id is None:
        st.info("Select a chat from the sidebar or create a new one.")